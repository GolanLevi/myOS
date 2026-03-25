import base64
import re
import io
from email.mime.text import MIMEText
from email.header import Header, decode_header, make_header
from utils.gmail_connector import get_gmail_service

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    print("beautifulsoup4 not installed - HTML emails will be read as raw text")

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("pdfplumber not installed - PDF attachments will not be readable")

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx not installed - DOCX attachments will not be readable")


def _extract_email_address(raw_value):
    match = re.search(r"<([^>]+)>", raw_value or "")
    if match:
        return match.group(1).strip().lower()
    if "@" in (raw_value or ""):
        return raw_value.strip().strip('"').lower()
    return ""


def _extract_display_name(raw_value):
    raw_value = (raw_value or "").strip()
    if not raw_value:
        return ""
    match = re.search(r"<([^>]+)>", raw_value)
    if match:
        return raw_value.replace(match.group(0), "").strip().strip('"')
    return raw_value.split("@")[0].strip().strip('"')


def _is_no_reply_address(email_value):
    email_value = (email_value or "").lower()
    return any(token in email_value for token in ["noreply", "no-reply", "notifications", "do-not-reply", "donotreply"])


def _decode_header_value(raw_value):
    value = (raw_value or "").strip()
    if not value:
        return ""
    try:
        return str(make_header(decode_header(value)))
    except Exception:
        return value


def _extract_body_from_parts(parts, mime_type="text/plain"):
    """מחלץ את גוף ההודעה רקורסיבית מתוך מבנה multipart"""
    body = ""
    for part in parts:
        if part.get("mimeType") == mime_type and part.get("body", {}).get("data"):
            data = part["body"]["data"]
            body += base64.urlsafe_b64decode(data).decode("utf-8", errors="replace")
        elif part.get("parts"):
            body += _extract_body_from_parts(part["parts"], mime_type)
    return body


def _html_to_text(html_content):
    """ממיר HTML לטקסט קריא"""
    if HAS_BS4:
        soup = BeautifulSoup(html_content, "html.parser")
        # הסרת script ו-style
        for tag in soup(["script", "style", "head"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    else:
        # Fallback — regex פשוט
        text = re.sub(r"<[^>]+>", "", html_content)
    
    # ניקוי שורות ריקות מיותרות
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _extract_attachments_info(parts):
    """מחלץ מידע על קבצים מצורפים (שם, סוג, גודל, מזהה)"""
    attachments = []
    for part in parts:
        filename = part.get("filename")
        if filename:
            size = part.get("body", {}).get("size", 0)
            mime = part.get("mimeType", "unknown")
            attachment_id = part.get("body", {}).get("attachmentId", "")
            attachments.append({
                "filename": filename,
                "mimeType": mime,
                "size": size,
                "attachmentId": attachment_id
            })
        if part.get("parts"):
            attachments.extend(_extract_attachments_info(part["parts"]))
    return attachments


# סוגי קבצים שאפשר לקרוא
READABLE_MIME_TYPES = [
    "application/pdf",
    "text/plain",
    "text/csv",
    "text/calendar",
    "application/ics",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # DOCX
    "application/msword",  # DOC
]

MAX_ATTACHMENT_SIZE = 5 * 1024 * 1024  # 5MB מקסימום


def _read_attachment_text(msg_id, attachment):
    """קורא תוכן קובץ מצורף בזיכרון (ללא הורדה לדיסק)"""
    filename = attachment.get("filename", "")
    mime = attachment.get("mimeType", "")
    att_id = attachment.get("attachmentId", "")
    size = attachment.get("size", 0)
    
    # בדיקות בסיסיות
    if not att_id:
        return None
    if mime not in READABLE_MIME_TYPES:
        return None
    if size > MAX_ATTACHMENT_SIZE:
        return f"[📎 {filename} — קובץ גדול מדי ({size // 1024}KB), לא נקרא]"
    
    try:
        service = get_gmail_service()
        att_data = service.users().messages().attachments().get(
            userId='me', messageId=msg_id, id=att_id
        ).execute()
        
        raw_bytes = base64.urlsafe_b64decode(att_data['data'])
        
        # --- PDF ---
        if mime == "application/pdf":
            if not HAS_PDF:
                return f"[📎 {filename} — PDF (לא ניתן לקרוא, חסר pdfplumber)]"
            
            pdf_file = io.BytesIO(raw_bytes)
            text_parts = []
            with pdfplumber.open(pdf_file) as pdf:
                for i, page in enumerate(pdf.pages[:10]):  # מקסימום 10 עמודים
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if text_parts:
                content = "\n".join(text_parts)
                # קיצור אם ארוך מדי
                if len(content) > 2000:
                    content = content[:2000] + "\n... [קוצר]"
                return f"📄 תוכן {filename}:\n{content}"
            return f"[📎 {filename} — PDF ריק או סרוק (ללא טקסט)]"
        
        # --- טקסט / CSV / ICS ---
        elif mime.startswith("text/") or mime == "application/ics":
            text = raw_bytes.decode("utf-8", errors="replace")
            if len(text) > 3000:
                text = text[:3000] + "\n... [קוצר]"
            return f"📄 תוכן {filename}:\n{text}"
        
        # --- DOCX ---
        elif mime in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            if not HAS_DOCX:
                return f"[📎 {filename} — DOCX (לא ניתן לקרוא, חסר python-docx)]"
            
            docx_file = io.BytesIO(raw_bytes)
            doc = docx.Document(docx_file)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # גם טבלאות
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            if text_parts:
                content = "\n".join(text_parts)
                if len(content) > 3000:
                    content = content[:3000] + "\n... [הטקסט קוצר]"
                return f"📄 תוכן {filename}:\n{content}"
            return f"[📎 {filename} — DOCX ריק (ללא טקסט)]"
        
    except Exception as e:
        print(f"Error reading attachment {filename}: {e}")
        return f"[📎 {filename} — שגיאה בקריאה]"
    
    return None


def get_full_email_body(message_data):
    """מחלץ את הגוף המלא של מייל — טקסט, HTML, קבצים מצורפים"""
    payload = message_data.get("payload", {})
    parts = payload.get("parts", [])
    
    # ניסיון 1: חילוץ text/plain
    body = ""
    if parts:
        body = _extract_body_from_parts(parts, "text/plain")
    
    # ניסיון 2: אם אין טקסט פשוט — ננסה HTML
    if not body.strip() and parts:
        html_body = _extract_body_from_parts(parts, "text/html")
        if html_body:
            body = _html_to_text(html_body)
    
    # ניסיון 3: אם אין parts — הגוף ישירות ב-payload.body
    if not body.strip() and payload.get("body", {}).get("data"):
        raw_data = payload["body"]["data"]
        decoded = base64.urlsafe_b64decode(raw_data).decode("utf-8", errors="replace")
        if payload.get("mimeType") == "text/html":
            body = _html_to_text(decoded)
        else:
            body = decoded
    
    # Fallback: Snippet
    if not body.strip():
        body = message_data.get("snippet", "")
    
    return body


def fetch_email_by_id(msg_id):
    """מושך מייל בודד עם כל התוכן המלא שלו"""
    service = get_gmail_service()
    try:
        msg_data = service.users().messages().get(userId='me', id=msg_id, format='full').execute()
        
        headers = msg_data.get("payload", {}).get("headers", [])
        subject = _decode_header_value(next((h['value'] for h in headers if h['name'] == 'Subject'), "No Subject"))
        sender = _decode_header_value(next((h['value'] for h in headers if h['name'] == 'From'), "Unknown"))
        reply_to = _decode_header_value(next((h['value'] for h in headers if h['name'] == 'Reply-To'), ""))
        to_value = _decode_header_value(next((h['value'] for h in headers if h['name'] == 'To'), ""))
        date = _decode_header_value(next((h['value'] for h in headers if h['name'] == 'Date'), ""))
        thread_id = msg_data.get("threadId", "")
        sender_email = _extract_email_address(sender)
        sender_name = _extract_display_name(sender)
        reply_to_email = _extract_email_address(reply_to)
        effective_reply_email = reply_to_email or sender_email
        reply_possible = bool(effective_reply_email) and not _is_no_reply_address(effective_reply_email)
        
        # גוף מלא
        full_body = get_full_email_body(msg_data)
        
        # קבצים מצורפים
        parts = msg_data.get("payload", {}).get("parts", [])
        attachments = _extract_attachments_info(parts) if parts else []
        
        # זיהוי הזמנות יומן (ICS)
        has_calendar_invite = any(
            a["mimeType"] in ["text/calendar", "application/ics"] for a in attachments
        )
        
        result = {
            "id": msg_id,
            "threadId": thread_id,
            "sender": sender,
            "sender_name": sender_name,
            "sender_email": sender_email,
            "reply_to": reply_to,
            "reply_to_email": reply_to_email,
            "effective_reply_email": effective_reply_email,
            "reply_possible": reply_possible,
            "recipient": to_value,
            "subject": subject,
            "date": date,
            "body": full_body,
            "snippet": msg_data.get("snippet", ""),
            "attachments": attachments,
            "has_calendar_invite": has_calendar_invite
        }
        
        # קריאת תוכן קבצים מצורפים (בזיכרון)
        if attachments:
            att_texts = []
            for att in attachments:
                content = _read_attachment_text(msg_id, att)
                if content:
                    att_texts.append(content)
                else:
                    att_texts.append(f"  📎 {att['filename']} ({att['mimeType']})")
            result["body"] += f"\n\n--- קבצים מצורפים ---\n" + "\n".join(att_texts)
        
        return result
    except Exception as e:
        print(f"Error fetching email {msg_id}: {e}")
        return None


def fetch_recent_emails(limit=5):
    """
    מושך את המיילים האחרונים מהתיבה (Inbox)
    מחזיר רשימה עם גוף מלא (לא רק snippet)
    """
    # Hard cap for stability and token savings
    if limit > 10:
        print(f"Limit {limit} is too high. Capping to 10 emails.")
        limit = 10

    service = get_gmail_service()
    
    print(f"Fetching last {limit} emails...")
    
    results = service.users().messages().list(userId='me', labelIds=['INBOX'], maxResults=limit).execute()
    messages = results.get('messages', [])

    clean_emails = []

    if not messages:
        print("No messages found.")
        return []

    for msg in messages:
        email_data = fetch_email_by_id(msg['id'])
        if email_data:
            clean_emails.append({
                "id": email_data["id"],
                "sender": email_data["sender"],
                "subject": email_data["subject"],
                "snippet": email_data["snippet"],
                "body": email_data["body"],
                "attachments": email_data["attachments"],
                "has_calendar_invite": email_data["has_calendar_invite"]
            })
    
    return clean_emails

def create_draft(to_email, subject, body, thread_id=None):
    """
    יוצר טיוטה חדשה ב-Gmail (לא שולח, רק שומר ב-Drafts)
    thread_id: אופציונלי - אם מסופק, הטיוטה תוגש כחלק מה-thread הקיים (כמו "Reply")
    """
    service = get_gmail_service()
    
    message = MIMEText(body, _charset="utf-8")
    message['to'] = to_email
    message['subject'] = Header(subject, "utf-8")
    
    raw = base64.urlsafe_b64encode(message.as_bytes()).decode()
    draft_body = {'message': {'raw': raw}}
    if thread_id:
        draft_body['message']['threadId'] = thread_id
    
    try:
        draft = service.users().drafts().create(userId='me', body=draft_body).execute()
        print(f"Draft created. ID: {draft['id']}")
        return draft
    except Exception as e:
        print(f"Error creating draft: {e}")
        return None

def send_email(to_email, subject, body, thread_id=None):
    """שולח מייל מיידית דרך ה-API
    thread_id: אופציונלי - מאפשר לשלוח כחלק מ-thread קיים (כמו "Reply")
    """
    service = get_gmail_service()
    try:
        message = MIMEText(body, _charset="utf-8")
        message['to'] = to_email
        message['subject'] = Header(subject, "utf-8")
        
        raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
        send_body = {'raw': raw_message}
        if thread_id:
            send_body['threadId'] = thread_id
        
        sent_message = service.users().messages().send(userId='me', body=send_body).execute()
        print(f"Email sent. ID: {sent_message['id']}")
        return sent_message['id']
        
    except Exception as e:
        print(f"Error sending email: {e}")
        return None



def create_label(label_name):
    """יוצר תווית חדשה בג'ימייל אם היא לא קיימת"""
    service = get_gmail_service()
    try:
        results = service.users().labels().list(userId='me').execute()
        labels = results.get('labels', [])
        for label in labels:
            if label['name'].lower() == label_name.lower():
                return label['id']
        
        label_object = {'name': label_name, 'labelListVisibility': 'labelShow', 'messageListVisibility': 'show'}
        created = service.users().labels().create(userId='me', body=label_object).execute()
        print(f"Label created: {label_name}")
        return created['id']
    except Exception as e:
        print(f"Error creating label: {e}")
        return None

def remove_label(msg_id, label_id_or_name):
    """מסיר תווית ממייל ספציפי לפי שם או ID (שימושי ל-mark_as_read וכו')"""
    service = get_gmail_service()
    # Gmail built-in labels like UNREAD, INBOX can be used directly by their string name as ID
    try:
        body = {'removeLabelIds': [label_id_or_name]}
        service.users().messages().modify(userId='me', id=msg_id, body=body).execute()
        print(f"Label '{label_id_or_name}' removed from email {msg_id}")
    except Exception as e:
        print(f"Error removing label: {e}")

def add_label_to_email(msg_id, label_name):
    """מוסיף תווית למייל ספציפי"""
    service = get_gmail_service()
    label_id = create_label(label_name)
    
    if label_id:
        try:
            body = {'addLabelIds': [label_id]}
            service.users().messages().modify(userId='me', id=msg_id, body=body).execute()
            print(f"Label '{label_name}' added to email {msg_id}")
        except Exception as e:
            print(f"Error adding label: {e}")

def archive_email(msg_id):
    """מעביר מייל לארכיון (מסיר אותו מה-Inbox)"""
    service = get_gmail_service()
    try:
        body = {'removeLabelIds': ['INBOX']}
        service.users().messages().modify(userId='me', id=msg_id, body=body).execute()
        print(f"Email {msg_id} archived.")
    except Exception as e:
        print(f"Error archiving email: {e}")

def trash_email(msg_id):
    """מעביר מייל לאשפה (Trash)"""
    service = get_gmail_service()
    try:
        service.users().messages().trash(userId='me', id=msg_id).execute()
        print(f"Email {msg_id} moved to TRASH.")
        return True
    except Exception as e:
        print(f"Error trashing email: {e}")
        return False
